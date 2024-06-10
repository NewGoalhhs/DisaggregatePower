from pathlib import Path

from docx import Document
from docx.shared import Mm

from app import __ROOT__


class ResultDocumentationHelper:
    def __init__(self, model_name):
        self.model_name = model_name
        # check if the document exists
        if Path(self.doc_path_model()).exists():
            self.document = Document(self.doc_path_model())
        else:
            self.document = Document()

    def doc_path(self):
        return __ROOT__ + '/docs'

    def doc_path_model(self):
        return self.doc_path() + '/' + self.model_name + '.docx'

    def add_image(self, image_path):
        self.document.add_picture(image_path, width=self.get_text_width())

    def add_heading(self, text, level):
        self.document.add_heading(text, level)

    def add_parameter(self, name, value):
        self.document.add_paragraph(name + ': ' + str(value))

    def add_break(self):
        self.document.add_page_break()

    def save(self):
        try:
            self.document.save(self.doc_path_model())
        except PermissionError:
            result = input("Permission denied. Please close the document.")
            if result == 'exit':
                return
            self.save()

    def get_text_width(self):
        section = self.document.sections[0]
        return Mm((section.page_width - section.left_margin - section.right_margin) / 36000)